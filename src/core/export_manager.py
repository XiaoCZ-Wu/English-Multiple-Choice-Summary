"""
导出管理模块
负责将题目导出为各种格式
"""

import csv
import os
from typing import List
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx2pdf import convert

from src.models import Question
from src.utils import CLASSIFICATIONS, get_timestamp, app_logger


class ExportOptions:
    """导出选项"""
    def __init__(
        self,
        include_answer: bool = False,
        include_answer_card: bool = False,
        include_source: bool = False
    ):
        self.include_answer = include_answer
        self.include_answer_card = include_answer_card
        self.include_source = include_source


class ExportManager:
    """导出管理器"""

    @staticmethod
    def export_to_docx(
        questions: List[Question],
        filepath: str,
        title: str = "",
        options: ExportOptions = None
    ) -> bool:
        """导出为Word文档"""
        try:
            if options is None:
                options = ExportOptions()
            
            doc = Document()
            
            # 设置样式
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            # 设置页边距
            section = doc.sections[0]
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            
            # 添加标题
            if title:
                p0 = doc.add_paragraph()
                run = p0.add_run(title)
                run.font.bold = True
                run.font.size = Pt(14)
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加题目
            answers = []
            for i, q in enumerate(questions):
                # 题目内容
                source_text = f"【{q.source}】" if options.include_source and q.source else ""
                p1 = doc.add_paragraph(f"{q.question}{source_text}", style="List Number")
                
                # 选项
                for opt, text in [("A", q.A), ("B", q.B), ("C", q.C), ("D", q.D)]:
                    p = doc.add_paragraph(f"{opt}. {text}")
                    p.paragraph_format.first_line_indent = Cm(0.63)
                
                answers.append(q.answer)
            
            # 添加答案
            if options.include_answer and answers:
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
                ans_str = ""
                for idx, a in enumerate(answers):
                    if idx != 0 and idx % 5 == 0:
                        ans_str += "\n"
                    ans_str += f"{idx + 1}. {a}\t\t"
                doc.add_paragraph(f"答案：\n{ans_str}")
            
            # 保存
            doc.save(filepath)
            return True
        except Exception as e:
            app_logger.error(f"导出DOCX失败: {e}")
            return False

    @staticmethod
    def export_to_pdf(
        questions: List[Question],
        filepath: str,
        title: str = "",
        options: ExportOptions = None
    ) -> bool:
        """导出为PDF"""
        try:
            # 先导出为DOCX，再转换
            temp_docx = filepath.replace(".pdf", "_temp.docx")
            if ExportManager.export_to_docx(questions, temp_docx, title, options):
                convert(temp_docx, filepath)
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                return True
            return False
        except Exception as e:
            app_logger.error(f"导出PDF失败: {e}")
            return False

    @staticmethod
    def export_to_csv(
        questions: List[Question],
        filepath: str
    ) -> bool:
        """导出为CSV"""
        try:
            with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow([
                    "序号", "题目", "选项A", "选项B", "选项C", "选项D",
                    "正确选项", "分类", "总计回答次数", "正答次数",
                    "套题名称", "分析"
                ])
                
                for idx, q in enumerate(questions):
                    writer.writerow([
                        idx + 1,
                        q.question,
                        q.A,
                        q.B,
                        q.C,
                        q.D,
                        q.answer,
                        CLASSIFICATIONS[q.classification] if 0 <= q.classification < len(CLASSIFICATIONS) else "Error",
                        q.total,
                        q.correct,
                        q.source,
                        q.analysis
                    ])
            return True
        except Exception as e:
            app_logger.error(f"导出CSV失败: {e}")
            return False
