import json
import os
import sys
import time
import zipfile
import csv
import random
import threading

from PySide6.QtCore import *
from PySide6.QtGui import *
from PySide6.QtUiTools import QUiLoader
from PySide6.QtWidgets import *
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
from docx2pdf import convert
from datetime import datetime
from tabulate import tabulate


class main(QObject):
    def __init__(self):
        super().__init__()
        with open("./data/questions.json", "r", encoding="utf-8") as file:
            self.questions = json.load(file)  # 所有题目
        print(self.questions)
        with open("./data/config.json", "r", encoding="utf-8") as file:
            self.config = json.load(file)  # 基本配置
        self.app = QApplication.instance()  # 当前应用实例
        self.main_window = None  # 主窗口
        self.window_title = None  # 记录窗口标题
        self.btn_group_1 = None  # 第一组RadioButton的按钮组
        self.btn_group_2 = None  # 第二组RadioButton的按钮组
        self.btn_group_3 = None  # 导出页面筛选条件的按钮组
        self.btn_group_4 = None  # 首页上练习模式的按钮组
        self.btn_group_5 = None  # 练习页面上的选项按钮
        self.papers = []  # 用于储存所有套题名称
        self.result_list = []  # 导出页面中tableWidget中被筛选的结果
        self.buttons = ()  # 所有的按钮
        self.functions = ()  # 与按钮链接的方法
        self.menu_1 = QMenu()  # 题目管理页面的右键菜单
        self.selected_rows = []  # 右键菜单选择的行编号
        self.deleted_questions = []  # 将删除但没保存的问题储存起来
        self.editable_rows = []  # 处于可编辑模式下的行
        self.edit_mode = False  # 编辑模式
        self.mode = -1  # 当前的练习模式（-1 - 没有任何模式；0 - 无尽；1 - 套题）
        self.prepared_questions = []  # 练习页面准备好的题目
        self.current_question = -1  # 当前正在练习页面上显示的问题
        self.single_time = 0  # 单题用时，单位：s
        self.total_time = 0  # 总用时，单位：s
        self.time_thread = None  # 多线程记录时间
        self.time_thread_event = None  # 线程事件
        self.questions_statistics = []  # 统计题目信息，下标0：0-错 1-对；下标1：分类的id；下标2：用时（s）
        self.show_answer = None  # 记录当前是否已经显示了答案

        # 设置字体（part1）
        self.font = QFont(self.config["font-name"], self.config["font-size"])
        self.app.setFont(self.font)

        self.build()

    def build(self):
        self.main_window = QUiLoader().load("./ui/main.ui")
        self.window_title = self.main_window.windowTitle()  # 记录标题
        self.main_window.installEventFilter(self)  # 给主窗口添加事件过滤器

        # 设置字体（part2）
        for textEdit in [
            self.main_window.textEdit, self.main_window.textEdit_2, self.main_window.textEdit_3,
            self.main_window.textEdit_4, self.main_window.textEdit_5, self.main_window.textEdit_6,
            self.main_window.textEdit_7, self.main_window.textEdit_8, self.main_window.textEdit_9,
            self.main_window.textEdit_10, self.main_window.textEdit_11
        ]:
            textEdit.setFont(self.font)

        # 按钮点击事件
        self.buttons = (  # 所有的按钮
            self.main_window.pushButton, self.main_window.pushButton_2, self.main_window.pushButton_3,
            self.main_window.pushButton_4, self.main_window.pushButton_5, self.main_window.pushButton_6,
            self.main_window.pushButton_7, self.main_window.pushButton_8, self.main_window.pushButton_9,
            self.main_window.pushButton_10, self.main_window.pushButton_11, self.main_window.pushButton_12,
            self.main_window.pushButton_13, self.main_window.pushButton_14, self.main_window.pushButton_15,
            self.main_window.pushButton_16, self.main_window.pushButton_17, self.main_window.pushButton_18,
            self.main_window.pushButton_19, self.main_window.pushButton_20, self.main_window.pushButton_21,
            self.main_window.pushButton_22
        )
        self.functions = (  # 与按钮链接的方法
            self.start, self.create, self.manage, self.settings, self.collect, self.back, self.reload_json, self.back,
            # “开始练习”按钮; “录入错题”按钮; “导出错题”按钮; “设置”按钮; 保存表单; 回到首页; 重新加载数据; 回到首页
            self.filter, self.select_all, self.export_format, self.rs_accu_rate, self.back, self.apply_settings,
            # 筛选题目; 导出错题页面筛选结果的全选; 导出错题; 重置筛选条件：正确率; 回到首页; 回到首页; self.select_dir
            self.select_dir, self.save_edits, self.backup, self.back, self.last_question, self.confirm_answer,
            # 选择导出文件的目录; 保存修改后的题目; 备份json数据; 返回首页（结束线程）; 切换到上一题; 提交答案并切换到下一题
            self.chat_robot, self.back
            # 将问题发送给AI; 返回首页（结束线程）
        )
        for idx, button in enumerate(self.buttons):
            button.clicked.connect(self.functions[idx])
        # 右键菜单
        self.init_actions()
        self.main_window.tableWidget.setContextMenuPolicy(Qt.CustomContextMenu)
        self.main_window.tableWidget.customContextMenuRequested.connect(self.show_right_click_menu)
        # 初始化
        self.init_settings_page()
        self.init_radio_groups()
        self.init_questions_list()

    def init_actions(self):
        """初始化右键菜单"""
        edit_action = QAction("编辑", self.main_window.tableWidget)
        edit_action.triggered.connect(self.edit_question)
        delete_action = QAction("删除", self.main_window.tableWidget)
        delete_action.triggered.connect(self.delete_question)
        self.menu_1.addAction(edit_action)
        self.menu_1.addAction(delete_action)

    def init_settings_page(self):
        """将读取到的配置文件应用到设置"""
        self.main_window.spinBox.setValue(self.config["font-size"])
        self.main_window.fontComboBox.setCurrentText(self.config["font-name"])
        self.main_window.lineEdit.setText(self.config["output_dir"])

    def init_radio_groups(self):
        """初始化RadioButton分组，实现组内互斥、组间独立"""
        # 定义两组RadioButton
        radioBtn_1 = [
            self.main_window.radioButton, self.main_window.radioButton_2,
            self.main_window.radioButton_3, self.main_window.radioButton_4
        ]
        radioBtn_2 = [
            self.main_window.radioButton_5, self.main_window.radioButton_6, self.main_window.radioButton_7,
            self.main_window.radioButton_8, self.main_window.radioButton_9, self.main_window.radioButton_10,
            self.main_window.radioButton_11, self.main_window.radioButton_12
        ]
        radioBtn_3 = [
            self.main_window.radioButton_13, self.main_window.radioButton_14, self.main_window.radioButton_15
        ]
        radioBtn_4 = [
            self.main_window.radioButton_16, self.main_window.radioButton_17
        ]
        radioBtn_5 = [
            self.main_window.radioButton_18, self.main_window.radioButton_19, self.main_window.radioButton_20,
            self.main_window.radioButton_21
        ]

        # 创建独立的按钮组，父控件设为主窗口
        self.btn_group_1 = QButtonGroup(self.main_window)
        self.btn_group_2 = QButtonGroup(self.main_window)
        self.btn_group_3 = QButtonGroup(self.main_window)
        self.btn_group_4 = QButtonGroup(self.main_window)
        self.btn_group_5 = QButtonGroup(self.main_window)

        # 将RadioButton加入对应组，并分配ID
        for idx, radioBtn in enumerate(radioBtn_1): self.btn_group_1.addButton(radioBtn, idx)
        for idx, radioBtn in enumerate(radioBtn_2): self.btn_group_2.addButton(radioBtn, idx)
        for idx, radioBtn in enumerate(radioBtn_3): self.btn_group_3.addButton(radioBtn, idx)
        for idx, radioBtn in enumerate(radioBtn_4): self.btn_group_4.addButton(radioBtn, idx)
        for idx, radioBtn in enumerate(radioBtn_5): self.btn_group_5.addButton(radioBtn, idx)

    def init_questions_list(self):
        """初始化题目管理页面"""
        # 向滚动区域添加错题
        self.add_question(questions=self.questions.copy(), is_init=True)
        # 将所有套题添加到下拉菜单
        self.main_window.comboBox.addItem("Any")
        self.main_window.comboBox.addItems(self.papers)

    def add_question(self, questions, is_init=False):
        """向导出错题页面的滚动区域添加错题"""
        def clear_rows():
            for idx in range(self.main_window.tableWidget.rowCount()):
                self.main_window.tableWidget.removeRow(0)

        # 准备工作
        clear_rows()
        keys = ["question", "A", "B", "C", "D", "answer", "classification", "source", "", "analysis"]
        classification = [
            "交际用语", "词义辨析", "时态", "非谓语动词", "定语从句", "状语从句", "情态动词", "名词性从句", "Error"
        ]
        self.main_window.tableWidget.setRowCount(len(questions))
        # 设置列宽
        self.main_window.tableWidget.setColumnWidth(0, self.config["font-size"] * 100)
        for col in range(1, 5): self.main_window.tableWidget.setColumnWidth(col, self.config["font-size"] * 30)
        self.main_window.tableWidget.setColumnWidth(8, self.config["font-size"] * 30)
        self.main_window.tableWidget.setColumnWidth(9, self.config["font-size"] * 50)
        # 添加数据
        for row, question in enumerate(questions):
            # 如果是初始化
            if (is_init is True) and (question["source"] != "") and (question["source"] not in self.papers):
                self.papers.append(f"{question['source']}")
                # # 如果直接点击“题目管理”-“全选”-“导出”就会报错显示result_list中啥都没有，勿删此行
                # self.result_list = self.questions.copy()
            for column in range(0, len(keys)):
                if column == 6:  # 分类
                    self.main_window.tableWidget.setItem(
                        row, column, QTableWidgetItem(classification[question.get(keys[column], 8)])
                    )
                    self.main_window.tableWidget.item(row, column).setTextAlignment(Qt.AlignCenter)
                    self.main_window.tableWidget.item(row, column).setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                elif column == 8:  # 正确率
                    try:
                        res = question["correct"] / question["total"]
                    except ZeroDivisionError:
                        res = question["correct"] / 1.00
                    except NameError:
                        QMessageBox.critical(self.main_window, "Error", "NameError in computing accuracy rate!")
                        res = "Error"
                    if isinstance(res, (int, float)): res = f"{res * 100:.2f}%"
                    self.main_window.tableWidget.setItem(row, 7, QTableWidgetItem(f"{res}"))
                    self.main_window.tableWidget.item(row, 7).setTextAlignment(Qt.AlignCenter)
                    self.main_window.tableWidget.item(row, 7).setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                elif column == 7:  # 套题名称
                    self.main_window.tableWidget.setItem(
                        row, column + 1, QTableWidgetItem(f"{question.get(keys[column], 'Error')}")
                    )
                    self.main_window.tableWidget.item(row, column + 1).setFlags(
                        Qt.ItemIsSelectable | Qt.ItemIsEnabled
                    )
                else:
                    self.main_window.tableWidget.setItem(
                        row, column, QTableWidgetItem(f"{question.get(keys[column], 'Error')}")
                    )
                    self.main_window.tableWidget.item(row, column).setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                if column == 5: self.main_window.tableWidget.item(row, 5).setTextAlignment(Qt.AlignCenter)

    def show_right_click_menu(self, position):
        """配置右键菜单"""
        # 1. 检查点击位置是否有效（可选，但推荐）
        clicked_item = self.main_window.tableWidget.itemAt(position)
        if not clicked_item: return
        # 2. 关键：获取当前所有被选中的行号
        # 使用 selectedItems() 并通过集合去重，这是最可靠的方法
        selected_rows = set()
        for item in self.main_window.tableWidget.selectedItems():
            selected_rows.add(item.row())
        # 将集合转换为排序后的列表，方便处理
        self.selected_rows = sorted(list(selected_rows))
        # 3. 如果没有任何行被选中（例如用户只是单击了一下），则默认选中当前点击的行
        if not self.selected_rows:
            self.selected_rows = [clicked_item.row()]
            self.main_window.tableWidget.selectRow(clicked_item.row())
        # 4. 显示菜单
        self.menu_1.exec(self.main_window.tableWidget.mapToGlobal(position))

    def start(self):
        """切换到开始练习页面，并开始练习"""
        def clear():  # 清空原有界面上的文字
            self.main_window.textEdit_9.clear()
            self.main_window.textEdit_10.clear()
            self.btn_group_5.setExclusive(False)
            for btn in self.btn_group_5.buttons(): btn.setChecked(False)
            self.btn_group_5.setExclusive(True)
            label = [
                self.main_window.label_14, self.main_window.label_15, self.main_window.label_16,
                self.main_window.label_17, self.main_window.label_18, self.main_window.label_19,
                self.main_window.label_20, self.main_window.label_21,
                self.main_window.radioButton_18, self.main_window.radioButton_19, self.main_window.radioButton_20,
                self.main_window.radioButton_21
            ]
            text = [
                "第 - 题，共 - 题（0.00%）", "单题用时：00: 00", "累计用时：00: 00", "第 - 次刷到该题", "过往正确率：0.00%",
                "当前练习模式：-", "所属套题：-", "正确答案：-", "A. 选项", "B. 选项", "C. 选项", "D. 选项"
            ]
            for idx, lbl in enumerate(label): lbl.setText(text[idx])
            # 初始化时间
            self.single_time = 0
            self.total_time = 0

        self.mode = self.btn_group_4.checkedId()  # 获取当前的练习模式
        if self.mode == -1: QMessageBox.information(self.main_window, "提示", "你还没有选择练习模式!")
        else:
            clear()  # 清空原先的内容
            self.main_window.stackedWidget.setCurrentIndex(4)
            # 准备题目
            self.prepared_questions = []  # 清空已准备的题目
            if self.mode == 0:  # 无尽模式
                # 创建题目列表（随机）
                self.prepared_questions = random.sample(self.questions[: 5], 5)
                self.update()
                self.time_thread = threading.Thread(target=self.record_time)
                self.time_thread.start()  # 计时开始
                self.time_thread_event = threading.Event()
            elif self.mode == 1:  # 套题模式
                pass

    def update(self):
        """切换题目"""
        # 题目编号 + 1
        self.current_question += 1
        # 设置按钮文本
        self.main_window.pushButton_20.setText("提交")
        # 设置状态
        self.show_answer = False
        # 重新计算单题时间
        self.single_time = 0
        # 设置选项状态
        self.btn_group_5.setExclusive(False)
        for btn in self.btn_group_5.buttons(): btn.setChecked(False)
        self.btn_group_5.setExclusive(True)
        # 更新题目
        self.main_window.textEdit_9.setText(
            self.prepared_questions[self.current_question].get("question", "Error!")
        )
        self.main_window.radioButton_18.setText(
            self.prepared_questions[self.current_question].get("A", "Error!")
        )
        self.main_window.radioButton_19.setText(
            self.prepared_questions[self.current_question].get("B", "Error!")
        )
        self.main_window.radioButton_20.setText(
            self.prepared_questions[self.current_question].get("C", "Error!")
        )
        self.main_window.radioButton_21.setText(
            self.prepared_questions[self.current_question].get("D", "Error!")
        )
        self.main_window.label_14.setText(
            f"第 {self.current_question + 1} 题，"
            f"共 {len(self.prepared_questions)} 题"
            f"（{((self.current_question + 1) / len(self.prepared_questions)) * 100:.2f}%）"
        )
        # self.main_window.label_15 单题用时在多线程内进行设置
        # self.main_window.label_16 累计用时在多线程内进行设置
        self.main_window.label_17.setText(
            f"第 {self.prepared_questions[self.current_question].get('total', -1000000) + 1} 次刷到该题"
        )
        try:
            res = self.prepared_questions[self.current_question]["correct"] \
                  / self.prepared_questions[self.current_question]["total"]
        except ZeroDivisionError:
            res = 0
        self.main_window.label_18.setText(
            f"过往正确率：{res * 100:.2f}%"
        )
        self.main_window.label_19.setText(
            f"当前练习模式：{['无尽模式', '套题模式', 'Error!', 'Error!'][self.mode]}"
        )
        self.main_window.label_20.setText(
            f"所属套题：{self.prepared_questions[self.current_question].get('source', 'Error!')}"
        )
        self.main_window.label_21.setText(
            f"正确答案：-"  # 点击确认答案按钮后才显示答案
        )
        self.main_window.label_23.setText(
            f"这题选什么呀？"  # 点击确认答案按钮后才显示答案
        )
        self.main_window.textEdit_10.setText(
            f""  # 点击确认答案按钮后才显示解析
        )

    def last_question(self):
        """点击并切换到上一道题"""
        pass

    def confirm_answer(self):
        """点击确认答案按钮后核对答案并显示解析"""
        # 判断按钮文本是否为“生成报告”
        if self.current_question == len(self.prepared_questions):
            self.main_window.stackedWidget.setCurrentIndex(5)
            self.show_report()
            return
        # 判断按钮文本是否为“下一题”
        if self.show_answer is True:  # 判断答案是否已经显示，如果显示了就下一题，没显示就核对答案
            self.update()
            return
        else:  # 核对答案
            # 获取当前的答案
            answer = self.btn_group_5.checkedId()
            if answer == -1:
                QMessageBox.information(self.main_window, "提示", "你没有选择任何选项!")
                return
            else:
                self.main_window.label_21.setText(
                    f"正确答案：{self.prepared_questions[self.current_question]['answer']}"
                )
                self.main_window.textEdit_10.setText(
                    f"{self.prepared_questions[self.current_question].get('analysis', 'Error!')}"
                )
                # 回答正确
                if ["A", "B", "C", "D"][answer] == self.prepared_questions[self.current_question]["answer"]:
                    self.main_window.label_23.setText(
                        f"正确!\n后面同学!"
                    )
                    self.questions_statistics.append(  # 记录正确率 分类 用时
                        [1, self.prepared_questions[self.current_question]["classification"], f"{self.single_time}"]
                    )
                # 回答错误
                else:
                    self.main_window.label_23.setText(
                        f"错误!\n都白讲了!\n来抬头我再说一遍→"
                    )
                    self.questions_statistics.append(  # 记录正确与否 分类 用时
                        [0, self.prepared_questions[self.current_question]["classification"], f"{self.single_time}"]
                    )
                print(self.questions_statistics)
            # 将做过的题合并到self.questions中
            for question in self.questions:
                if self.prepared_questions[self.current_question]["question"] == question["question"]:
                    question["correct"] += self.questions_statistics[-1][0]
                    question["total"] += 1
        # 设置按钮文本
        if self.current_question + 1 == len(self.prepared_questions):
            self.show_answer = False
            self.main_window.pushButton_20.setText("生成报告")
            # 让当前问题下标（从0计） == 问题数量，+= 1 后不存在与对应下标匹配的问题（相当于多出来一个）
            self.current_question += 1
        else:
            self.show_answer = True
            self.main_window.pushButton_20.setText("下一题")

    def show_report(self):
        """生成报告"""
        # 统计
        tmp_seconds = 0
        tmp_accu = 0
        classification_accu = [  # 类别 正确次数 总次数 正确率
            ["交际用语", 0, 0], ["词义辨析", 0, 0], ["时态", 0, 0], ["非谓语动词", 0, 0],
            ["定语从句", 0, 0], ["状语从句", 0, 0], ["情态动词", 0, 0], ["名词性从句", 0, 0]
        ]
        for each_question in self.questions_statistics:  # 下标0：0-错 1-对；下标1：分类的id；下标2：用时（s）
            # 分类正确率
            if each_question[0]:
                classification_accu[each_question[1]][1] += 1
            classification_accu[each_question[1]][2] += 1
            # 每道题用时的秒数之和
            tmp_seconds += int(each_question[2])
            # 统计一共对了多少道题
            tmp_accu += each_question[0]
        for idx, cls in enumerate(classification_accu):
            # 计算正确率
            try:
                classification_accu[idx].append(
                    f"{classification_accu[idx][1] / classification_accu[idx][2] * 100: .2f}%"
                )
            except ZeroDivisionError:
                classification_accu[idx].append("0.00%")

        total_time = f"{int(self.total_time / 60)}m{self.total_time % 60}s"
        average_seconds = int(tmp_seconds / len(self.questions_statistics))
        average_time = f"{int(average_seconds / 60)}m{average_seconds % 60}s"
        question_count = len(self.questions_statistics)
        for idx, item in enumerate([total_time, average_time, question_count]):
            self.main_window.textEdit_11.insertPlainText(
                f"{['累计用时', '平均用时', '题目总数'][idx]}".center(20)
            )
            self.main_window.textEdit_11.insertPlainText(f"{item}\n")
        self.main_window.textEdit_11.insertPlainText("=" * 80)
        self.main_window.textEdit_11.insertPlainText("\n")
        self.main_window.textEdit_11.insertPlainText(
            tabulate(classification_accu, ["题目类型", "答对", "共计", "正确率"], "grid")
        )
        # 将correct和total进行储存
        try:
            with open("./data/questions.json", "w", encoding="utf-8") as file:
                json.dump(self.questions, file, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"{e}")

    def record_time(self):
        """用于记录并修改单题、累计的时间"""
        while self.current_question != -1:
            time.sleep(1)
            # 单题用时
            self.single_time += 1
            minute_1 = int(self.single_time / 60)
            second_1 = self.single_time % 60
            self.main_window.label_15.setText(f"单题用时：{str(minute_1).zfill(2)}: {str(second_1).zfill(2)}")
            # 累计用时
            self.total_time += 1
            minute_2 = int(self.total_time / 60)
            second_2 = self.total_time % 60
            self.main_window.label_16.setText(f"累计用时：{str(minute_2).zfill(2)}: {str(second_2).zfill(2)}")
        print("Thread exit...")

    def chat_robot(self):
        """问答机器人，向AI提问"""
        QMessageBox.information(self.main_window, "提示", "此功能还没做")

    def create(self):
        """录入错题"""
        self.main_window.stackedWidget.setCurrentIndex(1)

    def collect(self):
        """收集错题表单并保存"""
        # 先检查数据
        if self.btn_group_1.checkedId() == -1:
            QMessageBox.information(self.main_window, "提示", "正确答案未选择！")
            return
        if self.btn_group_2.checkedId() == -1:
            QMessageBox.information(self.main_window, "提示", "分类未选择！")
            return
        new_question = {
            "question": self.main_window.textEdit.toPlainText(),
            "A": self.main_window.textEdit_2.toPlainText(),
            "B": self.main_window.textEdit_3.toPlainText(),
            "C": self.main_window.textEdit_4.toPlainText(),
            "D": self.main_window.textEdit_5.toPlainText(),
            "answer": ["A", "B", "C", "D"][self.btn_group_1.checkedId()],
            "classification": self.btn_group_2.checkedId(),
            "source": self.main_window.textEdit_6.toPlainText(),
            "analysis": self.main_window.textEdit_7.toPlainText(),
            "total": 0,
            "correct": 0
        }
        print("new:", new_question)
        # 如果是新分类，那么要添加在self.main_window.comboBox中
        if self.main_window.textEdit_6.toPlainText() not in self.papers:
            self.papers.append(self.main_window.textEdit_6.toPlainText())
            self.main_window.comboBox.addItem(self.main_window.textEdit_6.toPlainText())
        # 保存新问题
        try:
            self.questions.append(new_question)
            with open("./data/questions.json", "w", encoding="utf-8") as file:
                json.dump(self.questions, file, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"{e}")
        # 清空并重置表单
        for text_edit in [
            self.main_window.textEdit, self.main_window.textEdit_2, self.main_window.textEdit_3,
            self.main_window.textEdit_4, self.main_window.textEdit_5, self.main_window.textEdit_6,
            self.main_window.textEdit_7
        ]: text_edit.setPlainText("")
        self.btn_group_1.setExclusive(False)
        self.btn_group_2.setExclusive(False)
        for btn in self.btn_group_1.buttons(): btn.setChecked(False)
        for btn in self.btn_group_2.buttons(): btn.setChecked(False)
        self.btn_group_1.setExclusive(True)
        self.btn_group_2.setExclusive(True)

    def back(self):
        """返回首页"""
        # 从“开始练习”页面返回的
        if self.current_question != -1:
            self.current_question = -1
            self.time_thread = None
            self.time_thread_event = None
            self.single_time = 0
            self.total_time = 0
            self.questions_statistics = []
            self.main_window.textEdit_11.clear()
        # 从“题目管理”页面返回的
        self.main_window.setWindowTitle(self.window_title)  # 重新设置窗口标题
        self.reload_json(from_code=True)  # 返回相当于不保存

        # 切换页面
        self.main_window.stackedWidget.setCurrentIndex(0)

    def reload_json(self, from_code=False):
        """重新加载json数据文件"""
        try:
            with open("./data/questions.json", "r", encoding="utf-8") as file:
                self.questions = json.load(file)
            print("Reloading...")
            print(self.questions)
            if not from_code:
                QMessageBox.information(self.main_window, "提示", "Done!")
            else:
                # 题目管理界面的返回按钮==不保存
                for question in self.deleted_questions:
                    self.result_list.insert(question[0], question[1])
                self.deleted_questions.clear()  # 清空临时储存的删除但是不保存的题目
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"{e}")

    def filter(self):
        """筛选题目"""
        # 获取筛选标准
        boxes = [
            self.main_window.checkBox_10, self.main_window.checkBox_11, self.main_window.checkBox_12,
            self.main_window.checkBox_13, self.main_window.checkBox_14, self.main_window.checkBox_15,
            self.main_window.checkBox_16, self.main_window.checkBox_17
        ]
        checkedId = self.btn_group_3.checkedId()
        papers = self.main_window.comboBox.currentText()
        classification = []
        for idx, checkBox in enumerate(boxes):
            if checkBox.isChecked() is True: classification.append(idx)
        # 进行筛选
        filtered_by_classification = set()
        filtered_by_checkedId = set()
        filtered_by_papers = set()
        for question in self.questions:
            if (classification != []) and (question["classification"] in classification):
                filtered_by_classification.add(tuple(question.items()))
            if question["source"] == papers:
                filtered_by_papers.add(tuple(question.items()))
            try:
                if (checkedId != -1) and (question["correct"] / question["total"] <= [0.25, 0.50, 0.75][checkedId]):
                    filtered_by_checkedId.add(tuple(question.items()))
            except ZeroDivisionError:
                filtered_by_checkedId.add(tuple(question.items()))
        # 集合运算筛选题目：一有两无全并集，两有一无两交集，三有零无全交集
        non_empty = [s for s in (filtered_by_classification, filtered_by_checkedId, filtered_by_papers) if s]
        if len(non_empty) != 0:
            if len(non_empty) == 1:
                result_set = filtered_by_classification | filtered_by_checkedId | filtered_by_papers
            elif len(non_empty) == 2:
                result_set = non_empty[0] & non_empty[1]
            else:
                result_set = filtered_by_classification & filtered_by_checkedId & filtered_by_papers
            self.result_list.clear()
            for each_tuple in result_set: self.result_list.append(dict(each_tuple))
            # 添加到页面中
            self.add_question(questions=self.result_list.copy(), is_init=False)
            self.main_window.setWindowTitle(f"{self.window_title} - {len(self.result_list)}个筛选结果")
        else:
            if papers == "Any":
                # 全不选等于全选
                self.result_list = self.questions.copy()
                self.add_question(questions=self.result_list.copy(), is_init=False)
                self.main_window.setWindowTitle(f"{self.window_title} - {len(self.questions)}个筛选结果")

    def rs_accu_rate(self):
        """重置按钮组：题目管理中筛选正确率的按钮组"""
        self.btn_group_3.setExclusive(False)
        for btn in self.btn_group_3.buttons(): btn.setChecked(False)
        self.btn_group_3.setExclusive(True)

    def manage(self):
        """题目管理页面"""
        self.main_window.stackedWidget.setCurrentIndex(2)
        self.filter()
        self.add_question(questions=self.result_list.copy(), is_init=False)

    def select_all(self):
        """全选已筛选的结果"""
        self.main_window.tableWidget.selectAll()

    def export_format(self):
        """选择导出的格式：.docx/.pdf"""
        # 如果存在未保存的情况
        if (len(self.deleted_questions) != 0) or (self.edit_mode is True):
            QMessageBox.information(self.main_window, "提示", "存在未保存的数据，请先保存再导出！")
            return
        # 获取已选择的筛选结果
        selected_questions = []
        for range_obj in self.main_window.tableWidget.selectedRanges():
            for row in range(range_obj.topRow(), range_obj.bottomRow() + 1):
                selected_questions.append(self.result_list[row])

        def export_docx(path, filename, title, from_pdf=False, additions=None):
            # 写入docx文件
            doc = Document()

            normal_style = doc.styles['Normal']  # 设置“正文”样式
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            section = doc.sections[0]  # 设置页边距
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

            if title != "":
                p0 = doc.add_paragraph()
                run = p0.add_run(title)
                run.font.bold = True
                run.font.size = Pt(14)
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

            ans = []  # 记录每道题的答案
            for question in selected_questions:
                if additions[2] is True: add = f"【{question.get('source', 'Error!')}】"  # 是否添加套题标注
                else: add = ""
                p1 = doc.add_paragraph(f"{question.get('question', 'Error!')}{add}", style="List Number")
                p2 = doc.add_paragraph(f"A. {question.get('A', 'Error!')}")
                p2.paragraph_format.first_line_indent = Cm(0.63)
                p3 = doc.add_paragraph(f"B. {question.get('B', 'Error!')}")
                p3.paragraph_format.first_line_indent = Cm(0.63)
                p4 = doc.add_paragraph(f"C. {question.get('C', 'Error!')}")
                p4.paragraph_format.first_line_indent = Cm(0.63)
                p5 = doc.add_paragraph(f"D. {question.get('D', 'Error!')}")
                p5.paragraph_format.first_line_indent = Cm(0.63)
                ans.append(question.get("answer", "?"))
                if question.get("answer", "?") == "?": print("\n", "=" * 20, f"\n{question}")
            if additions[0] is True: # 是否写入答案
                doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)  # 分页符
                # 写入答案
                ans_str = ""
                for idx, a in enumerate(ans):
                    if idx != 0 and idx % 5 == 0: ans_str += "\n"
                    ans_str += f"{idx + 1}. {a}\t\t"
                doc.add_paragraph(f"答案：\n{ans_str}")
                print(ans_str)
            # 是否添加答题卡


            doc.save(f"{path}{filename}.docx")
            if not from_pdf:
                QMessageBox.information(self.main_window, "提示", f"文档{filename}.docx创建成功！")

        def export_pdf(path, filename, title, additions):
            # 先写入docx再转为pdf
            export_docx(path=path, filename=filename, title=title, from_pdf=True, additions=additions)
            convert(f"{path}{filename}.docx", f"{path}{filename}.pdf")
            if os.path.exists(f"{path}{filename}.docx"): os.remove(f"{path}{filename}.docx")
            QMessageBox.information(self.main_window, "提示", f"文档{filename}.pdf创建成功！")

        def export_csv(path, filename):
            with open(f"{path}{filename}.csv", "w", newline="", encoding="utf-8-sig") as file:
                writer = csv.writer(file)
                writer.writerow(
                    [
                        "序号", "题目", "选项A", "选项B", "选项C", "选项D", "正确选项",
                        "分类", "总计回答次数", "正答次数", "套题名称", "分析"
                    ]
                )
                classification = [
                    "交际用语", "词义辨析", "时态", "非谓语动词", "定语从句", "状语从句", "情态动词", "名词性从句", "Error"
                ]
                for idx, question in enumerate(selected_questions):
                    writer.writerow(
                        [
                            idx + 1,
                            question.get("question", "Error!"),
                            question.get("A", "Error!"),
                            question.get("B", "Error!"),
                            question.get("C", "Error!"),
                            question.get("D", "Error!"),
                            question.get("answer", "Error!"),
                            classification[question.get("classification", 8)],
                            question.get("total", "Error!"),
                            question.get("correct", "Error!"),
                            question.get("source", "Error!"),
                            question.get("analysis", "Error!")
                        ]
                    )
            QMessageBox.information(self.main_window, "提示", f"文档{filename}.csv创建成功！")

        if len(selected_questions) == 0:
            QMessageBox.information(self.main_window, "提示", "你还没有选择任何题目!")
        else:
            def collect_and_generate():
                is_write_answer = False
                is_write_ans_card = False
                is_write_source = False
                if lineEdit_1.text() != "": filename = lineEdit_1.text()
                else: filename = "new"
                if lineEdit_2.text() != "": title = lineEdit_2.text()
                else: title = ""
                if lineEdit_3.text() != "":
                    if (lineEdit_3.text()[-1] != "/") or (lineEdit_3.text()[-1] != "\\"):
                        path = lineEdit_3.text() + "/"
                    else: path = lineEdit_3.text()
                else: path = "./output/"
                if checkBox_1.isChecked() is True: is_write_answer = True
                if checkBox_2.isChecked() is True: is_write_ans_card = True
                if checkBox_3.isChecked() is True: is_write_source = True
                # 如果没有预先安装Microsoft Office会报错
                try:
                    if btn_group.checkedId() == 0:
                        export_docx(
                            path=path, filename=filename, title=title, additions=[
                                is_write_answer, is_write_ans_card, is_write_source
                            ]
                        )
                    elif btn_group.checkedId() == 1:
                        export_pdf(
                            path=path, filename=filename, title=title, additions=[
                                is_write_answer, is_write_ans_card, is_write_source
                            ]
                        )
                    elif btn_group.checkedId() == 2:
                        export_csv(path=path, filename=filename)
                    else:
                        QMessageBox.information(self.main_window, "提示", "你还没有选择任何文件类型!")
                except Exception as e:
                    QMessageBox.critical(self.main_window, "Error", f"{e}")


            # 创建对话框收集生成文档的信息
            dialog = QDialog()
            dialog.setWindowTitle("导出文件")
            dialog.setModal(True)
            lineEdit_1 = QLineEdit()
            lineEdit_1.setPlaceholderText("请输入文件名，默认为New")
            lineEdit_1.setMinimumWidth(300)
            lineEdit_2 = QLineEdit()
            lineEdit_2.setPlaceholderText("请输入文档标题，可选")
            lineEdit_2.setMinimumWidth(300)
            lineEdit_3 = QLineEdit()
            lineEdit_3.setPlaceholderText("请指定生成目录，默认为./output/")
            lineEdit_3.setMinimumWidth(300)
            lineEdit_3.setText(self.config["output_dir"])
            checkBox_1 = QCheckBox("将答案写入文件末尾（答案单独一页）")
            checkBox_2 = QCheckBox("创建对应题库并附带答题卡")
            checkBox_3 = QCheckBox("在题目结尾添加出处")
            r_btn_docx = QRadioButton(".docx")
            r_btn_pdf = QRadioButton(".pdf")
            r_btn_csv = QRadioButton(".csv")
            btn_group = QButtonGroup(dialog)
            btn_group.addButton(r_btn_docx, 0)
            btn_group.addButton(r_btn_pdf, 1)
            btn_group.addButton(r_btn_csv, 2)
            buttons = QDialogButtonBox.Ok | QDialogButtonBox.Cancel
            button_box = QDialogButtonBox(buttons, dialog)
            button_box.accepted.connect(dialog.accept)
            button_box.rejected.connect(dialog.reject)
            layout = QFormLayout(dialog)
            layout.addRow("文件名称", lineEdit_1)
            layout.addRow("文档标题", lineEdit_2)
            layout.addRow("文件路径", lineEdit_3)
            layout.addWidget(r_btn_docx)
            layout.addWidget(r_btn_pdf)
            layout.addWidget(r_btn_csv)
            layout.addWidget(checkBox_1)
            layout.addWidget(checkBox_2)
            layout.addWidget(checkBox_3)
            layout.addWidget(button_box)
            dialog_result = dialog.exec()
            if dialog_result == QDialog.Accepted: collect_and_generate()

    def edit_question(self):
        """题目管理页面：右键菜单 - 编辑"""
        if not self.selected_rows:
            QMessageBox.information(self.main_window, "提示", "没有选中任何行可编辑。")
            return
        # if len(self.selected_rows) != 1:
        #     QMessageBox.information(self.main_window, "提示", "只有选择1行时才可编辑。")
        #     return
        for row in self.selected_rows:
            # 将该行的单元格设置为可编辑
            for col in range(self.main_window.tableWidget.columnCount()):
                if col == 7: continue  # 正确率禁止编辑
                item = self.main_window.tableWidget.item(row, col)
                if item:
                    item.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled | Qt.ItemIsEditable)
            # 记录可编辑的行
            self.editable_rows.append(row)
        # 开启编辑模式
        self.edit_mode = True

    def delete_question(self):
        """题目管理页面：右键菜单 - 删除选中的行"""
        if not self.selected_rows:
            QMessageBox.information(self.main_window, "提示", "没有选中任何行可删除。")
            return
        # 弹出确认对话框
        reply = QMessageBox.question(
            self.main_window, '确认删除', f"确定要删除选中的{len(self.selected_rows)}行吗？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if reply == QMessageBox.No: return
        # 从最后一行开始倒序删除以避免行号错乱
        for row in sorted(self.selected_rows, reverse=True):
            # 从self.questions中删除
            try:
                for idx, question in enumerate(self.questions):
                    if question == self.result_list[row]:
                        del self.questions[idx]
            except Exception as e:
                print(f"from self.question; {e}")
            # 从self.result_list中删除并保存防止用户点击“返回”（也就是不保存）
            # 将question重新添加到self.result_list，防止list index out of range
            try:
                self.deleted_questions.append([row, self.result_list.pop(row)])  # 记录下标，方便插入，保持原来顺序
            except Exception as e:
                print(f"from self.result_list; {e}")
            # 从界面self.main_window.tableWidget中删除
            self.main_window.tableWidget.removeRow(row)

        # 刷新表格显示
        self.add_question(questions=self.result_list.copy(), is_init=False)

    def save_edits(self):
        """点击题目管理页面的保存按钮后保存修改后的题目"""
        invalid = []  # 记录非法的类型
        is_message = False  # 非法数据提示框是否已经弹出
        def check(target, data):
            # 检查各种参数是否合法
            if target == "answer":
                if data in ["A", "B", "C", "D"]: return data
                elif "invalid-answer" not in invalid: invalid.append("invalid-answer")
            if target == "classification":
                for idx, classification in enumerate(
                    ["交际用语", "词义辨析", "时态", "非谓语动词", "定语从句", "状语从句", "情态动词", "名词性从句"]
                ):
                    if data == classification: return idx
                if "invalid-classification" not in invalid: invalid.append("invalid-classification")
            if target == "source":
                # 检查是否是新试卷
                if data in self.papers: return data
                else:
                    self.papers.append(data)
                    self.main_window.comboBox.addItem(data)
                    return data

        # 获取新修改的表单并修改self.question和self.result_list
        if len(self.editable_rows) != 0:
            for row in self.editable_rows:
                for idx, question in enumerate(self.questions):
                    if question["question"] == self.result_list[row]["question"]:
                        modified_question = {
                            "question": self.main_window.tableWidget.item(row, 0).text(),
                            "A": self.main_window.tableWidget.item(row, 1).text(),
                            "B": self.main_window.tableWidget.item(row, 2).text(),
                            "C": self.main_window.tableWidget.item(row, 3).text(),
                            "D": self.main_window.tableWidget.item(row, 4).text(),
                            "answer": check(
                                target="answer", data=self.main_window.tableWidget.item(row, 5).text()
                            ),
                            "classification": check(
                                target="classification", data=self.main_window.tableWidget.item(row, 6).text()
                            ),
                            "source": check(
                                target="source", data=self.main_window.tableWidget.item(row, 8).text()
                            ),
                            "analysis": self.main_window.tableWidget.item(row, 9).text(),
                            "total": self.result_list[row]["total"],
                            "correct": self.result_list[row]["correct"]
                        }
                        if len(invalid) != 0:
                            QMessageBox.critical(
                                self.main_window, "Error",
                                f"存在非法数据，请修改后再保存！\ninvalid ids: {invalid}\nrow: {row}"
                            )
                        else:
                            self.questions[idx] = modified_question.copy()
                            self.result_list[row] = modified_question.copy()
                            for column in range(0, 10):
                                self.main_window.tableWidget.item(row, column).setFlags(
                                    Qt.ItemIsSelectable | Qt.ItemIsEnabled
                                )
                            print(self.questions[idx])
        self.editable_rows.clear()  # 清除可编辑的行
        self.selected_rows.clear()  # 清除被选择的行
        self.deleted_questions.clear()  # 清除删除但是未保存的数据
        self.edit_mode = False  # 修改编辑模式
        try:
            with open("./data/questions.json", "w", encoding="utf-8") as file:
                json.dump(self.questions, file, ensure_ascii=False, indent=2)
            if len(invalid) == 0:
                self.check_comboBox()
                QMessageBox.information(self.main_window, "提示", "题库保存成功!")
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"{e}")

    def check_comboBox(self):
        """检查删除题目后，是否还存在该分类（实际上是重新添加了一遍分类）"""
        # 清空
        self.papers = []
        self.main_window.comboBox.clear()
        # 重新添加
        for question in self.questions:
            if question["source"] not in self.papers: self.papers.append(f"{question['source']}")
        self.main_window.comboBox.addItem("Any")
        self.main_window.comboBox.addItems(self.papers)

    def settings(self):
        """切换到设置页面"""
        self.main_window.stackedWidget.setCurrentIndex(3)

    def select_dir(self):
        """选择导出文件的目录"""
        selected_dir = QFileDialog.getExistingDirectory(self.main_window, "请选择一个目录")
        if selected_dir == "": selected_dir = "./output/"
        self.main_window.lineEdit.setText(selected_dir)

    def apply_settings(self):
        """将设置中的内容应用并保存配置文件"""
        # 收集信息
        self.config["font-size"] = self.main_window.spinBox.value()
        self.config["font-name"] = self.main_window.fontComboBox.currentText()
        self.config["output_dir"] = self.main_window.lineEdit.text()
        with open("./data/config.json", "w", encoding="utf-8") as file:
            json.dump(self.config, file, ensure_ascii=False, indent=2)
        # 重新启动应用设置
        try:
            os.execv(sys.executable, ['python'] + sys.argv)
        except:
            os.execv(sys.executable, sys.argv)

    def backup(self):
        """备份json数据"""
        try:
            t = datetime.now().strftime('%y%m%d_%H%M%S')
            with zipfile.ZipFile(
                    f"./backup/backup_{t}.zip",
                    "w",
                    zipfile.ZIP_DEFLATED
            ) as file:
                file.write("./data/questions.json", arcname="questions.json")
                file.write("./data/config.json", arcname="config.json")
            QMessageBox.information(self.main_window, "提示", f"json数据备份成功！\n./backup/backup_{t}.zip")
        except Exception as e:
            QMessageBox.critical(self.main_window, "Error", f"{e}")

    def eventFilter(self, obj, event):
        """事件过滤：关闭窗口（终极修复版）"""
        if obj == self.main_window and event.type() == QEvent.Close:
            # 1. 先忽略当前的关闭事件（阻止窗口直接关闭）
            event.ignore()
            # 2. 弹出确认框
            reply = QMessageBox.question(
                self.main_window, "提示",
                "是否确定退出? \n在退出前请确保所有数据均已保存，防止丢失!\n做题时必须点击“生成报告”按钮，否则数据不会保存!",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )
            # 3. 根据选择手动控制窗口
            if reply == QMessageBox.Yes:
                self.back()  # 执行清理逻辑
                self.main_window.close()  # 手动关闭窗口
            # 4. 无论选什么，都返回True表示事件已处理
            return True
        return super().eventFilter(obj, event)

if __name__ == '__main__':
    app = QApplication()
    Run = main()
    Run.main_window.show()
    sys.exit(app.exec())
